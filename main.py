<<<<<<< HEAD
#!/usr/bin/env python
# encoding: utf-8

from docx import document
# from docx import Document
# from docx.shared import Inches


def p(text):
    print(type(text))
    print(text)

# document = Document()
# p(document)
#
# document.add_heading('Document Title', 0)
#
# p = document.add_paragraph('A plain paragraph having some ')
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True
#
# document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='IntenseQuote')
#
# document.add_paragraph(
#     'first item in unordered list', style='ListBullet'
# )
# document.add_paragraph(
#     'first item in ordered list', style='ListNumber'
# )
#
#
# table = document.add_table(rows=1, cols=3)
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'Qty'
# hdr_cells[1].text = 'Id'
# hdr_cells[2].text = 'Desc'
# document.add_page_break()
#
# document.save('demo.docx')
=======
import urllib2

enable_proxy = True

proxy_handler = urllib2.ProxyHandler({"http": 'http://some-proxy.com:8080'})

null_proxy_handler = urllib2.ProxyHandler({})

if enable_proxy:
    opener = urllib2.build_opener(proxy_handler)
else:
    opener = urllib2.build_opener(null_proxy_handler)

urllib2.install_opener(opener)
>>>>>>> a642f9235fda4a432ad5afa0a896bc18e6dc1caf
