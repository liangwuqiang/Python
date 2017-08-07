# -*- coding: cp936 -*-
from sgmllib import SGMLParser
import os
import sys
import urllib
#import urllib2
from docx import Document
from docx.shared import Inches
import time


def testprint(txt):
    print(txt)
    print(type(txt))


# 获取要解析的url
class GetUrl(SGMLParser):
    def __init__(self):
        SGMLParser.__init__(self)
        self.start = False
        self.urlArr = []

    def start_div(self, attr):
        for name, value in attr:
            if value == "ChairmanCont Bureau":  # 页面js中的固定值
                self.start = True

    def end_div(self):
        self.start = False

    def start_a(self, attr):
        if self.start:
            for name, value in attr:
                self.urlArr.append(value)

    def getUrlArr(self):
        return self.urlArr


# 解析上面获取的url，获取有用数据
class getManInfo(SGMLParser):
    def __init__(self):
        SGMLParser.__init__(self)
        self.start = False
        self.p = False
        self.dl = False
        self.manInfo = []
        self.subInfo = []

    def start_div(self, attr):
        for name, value in attr:
            if value == "SpeakerInfo":  # 页面js中的固定值
                self.start = True

    def end_div(self):
        self.start = False

    def start_p(self, attr):
        if self.dl:
            self.p = True

    def end_p(self):
        self.p = False

    def start_img(self, attr):
        if self.dl:
            for name, value in attr:
                self.subInfo.append(value)

    def handle_data(self, data):
        if self.p:
            self.subInfo.append(data.decode('utf-8'))

    def start_dl(self, attr):
        if self.start:
            self.dl = True

    def end_dl(self):
        self.manInfo.append(self.subInfo)
        self.subInfo = []
        self.dl = False

    def getManInfo(self):
        return self.manInfo


# urlSource = "http://www.XXX"
# sourceData = urllib.urlopen(urlSource).read()
with open("new.html") as f:
    sourceData = f.read()
startTime = time.clock()
##get urls
getUrl = GetUrl()
getUrl.feed(sourceData)


urlArr = getUrl.getUrlArr()

testprint(urlArr)

getUrl.close()
# print("get url use:" + str((time.clock() - startTime)))
startTime = time.clock()
##get maninfos
manInfos = getManInfo()
# for url in urlArr:  # one url one person
#     data = urllib.urlopen(url).read()

with open("new.html") as f:
    data = f.read()

manInfos.feed(data)

infos = manInfos.getManInfo()
manInfos.close()
# print("get maninfos use:" + str((time.clock() - startTime)))
startTime = time.clock()
# word 文档
saveFile = os.getcwd() + "/xxx.docx"
doc = Document()
##word title 标题
# doc.add_heading("HEAD".decode('gbk'), 0)
doc.add_heading("HEAD", 0)
p = doc.add_paragraph("HEADCONTENT:")

##write info
for infoArr in infos:
    i = 0
    for info in infoArr:
        if i == 0:  ##img url
            arr1 = info.split('.')
            suffix = arr1[len(arr1) - 1]
            arr2 = info.split('/')
            preffix = arr2[len(arr2) - 2]
            imgFile = os.getcwd() + "\\imgs\\" + preffix + "." + suffix
            if not os.path.exists(os.getcwd() + "\\imgs"):
                os.mkdir(os.getcwd() + "\\imgs")
            imgData = urllib.urlopen(info).read()

            try:
                f = open(imgFile, 'wb')
                f.write(imgData)
                f.close()
                doc.add_picture(imgFile, width=Inches(1.25))
                os.remove(imgFile)
            except Exception as err:
                print (err)


        elif i == 1:
            doc.add_heading(info + ":", level=1)
        else:
            doc.add_paragraph(info, style='ListBullet')
        i = i + 1

doc.save(saveFile)
# print("word use:" + str((time.clock() - startTime)))

# temp = infos
# print(temp)
# print(type(temp))

